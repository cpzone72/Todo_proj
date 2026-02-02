import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name='Malgun Gothic', font_size=None, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = Pt(font_size)
    if bold:
        run.font.bold = True

def create_bmt_word_doc_v3():
    # Load data
    with open('bmt_data.json', 'r', encoding='utf-8') as f:
        data = json.load(f)

    document = Document()
    
    # Set default style
    style = document.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # Title
    title_paragraph = document.add_heading(level=0)
    title_run = title_paragraph.add_run(data['title'])
    set_font(title_run, font_size=20, bold=True)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Overview
    h1 = document.add_heading(level=1)
    set_font(h1.add_run('1. 개요 및 목적'), font_size=14, bold=True)
    
    p = document.add_paragraph()
    run = p.add_run(data['overview']['target_label'])
    set_font(run, bold=True)
    run = p.add_run(data['overview']['target_value'])
    set_font(run)
    
    run = p.add_run(data['overview']['purpose_label'])
    set_font(run, bold=True)
    run = p.add_run(data['overview']['purpose_value'])
    set_font(run)

    # 2. Environment
    h2 = document.add_heading(level=1)
    set_font(h2.add_run(data['environment']['title']), font_size=14, bold=True)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    
    p = hdr_cells[0].paragraphs[0]
    run = p.add_run('구분')
    set_font(run, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = hdr_cells[1].paragraphs[0]
    run = p.add_run('상세 내역')
    set_font(run, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for item, detail in data['environment']['items']:
        row_cells = table.add_row().cells
        set_font(row_cells[0].paragraphs[0].add_run(item))
        set_font(row_cells[1].paragraphs[0].add_run(detail))

    # 3. Detailed Procedures
    h3 = document.add_heading(level=1)
    set_font(h3.add_run(data['procedures_title']), font_size=14, bold=True)

    for section in data['sections']:
        h_sec = document.add_heading(level=2)
        set_font(h_sec.add_run(section['title']), font_size=12, bold=True)
        
        # Add table for tests
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['ID', '항목', '세부 항목', '테스트 절차 및 방법', '평가 기준']
        
        for i, header_text in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(header_text)
            set_font(run, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row_data in section['tests']:
            row_cells = table.add_row().cells
            for i, text in enumerate(row_data):
                set_font(row_cells[i].paragraphs[0].add_run(text))

    document.add_paragraph('\n')
    p_footer = document.add_paragraph()
    set_font(p_footer.add_run(data['footer']))

    # Save
    file_name = 'NVR3864-V6-IQ_BMT_Test_Procedures_Fixed.docx'
    document.save(file_name)
    print(f"Word document saved as {file_name}")

if __name__ == "__main__":
    create_bmt_word_doc_v3()
