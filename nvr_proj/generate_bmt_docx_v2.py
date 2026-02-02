import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_bmt_word_doc_v2():
    # Load data
    with open('bmt_data.json', 'r', encoding='utf-8') as f:
        data = json.load(f)

    document = Document()

    # Title
    title = document.add_heading(data['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Overview
    document.add_heading('1. 개요 및 목적', level=1)
    p = document.add_paragraph()
    run = p.add_run(data['overview']['target_label'])
    run.bold = True
    p.add_run(data['overview']['target_value'])
    
    run = p.add_run(data['overview']['purpose_label'])
    run.bold = True
    p.add_run(data['overview']['purpose_value'])

    # 2. Environment
    document.add_heading(data['environment']['title'], level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '구분'
    hdr_cells[1].text = '상세 내역'
    
    for item, detail in data['environment']['items']:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = detail

    # 3. Detailed Procedures
    document.add_heading(data['procedures_title'], level=1)

    for section in data['sections']:
        document.add_heading(section['title'], level=2)
        
        # Add table for tests
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = '항목'
        hdr_cells[2].text = '세부 항목'
        hdr_cells[3].text = '테스트 절차 및 방법'
        hdr_cells[4].text = '평가 기준'

        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row_data in section['tests']:
            row_cells = table.add_row().cells
            for i, text in enumerate(row_data):
                row_cells[i].text = text

    document.add_paragraph('\n')
    document.add_paragraph(data['footer'])

    # Save
    file_name = 'NVR3864-V6-IQ_BMT_Test_Procedures.docx'
    document.save(file_name)
    print(f"Word document saved as {file_name}")

if __name__ == "__main__":
    create_bmt_word_doc_v2()
